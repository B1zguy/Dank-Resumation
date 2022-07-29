import docx

def docx_replace(doc, data):
    paragraphs = list(doc.paragraphs)
    print('start')
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)

    def replace_inline(key_name, val, inline):
        # Replace strings and retain the same style.
        # The text to be replaced can be split over several runs so
        # search through, identify which runs need to have text replaced
        # then replace the text in those identified
        started = False
        key_index = 0
        # found_runs is a list of (inline index, index of match, length of match)
        found_runs = list()
        found_all = False
        replace_done = False
        for i in range(len(inline)):

            # case 1: found in single run so short circuit the replace
            if key_name in inline[i].text and not started:
                found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                text = inline[i].text.replace(key_name, str(val))
                inline[i].text = text
                replace_done = True
                found_all = True
                break

            if key_name[key_index] not in inline[i].text and not started:
                # keep looking ...
                continue

            # case 2: search for partial text, find first run
            if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                # check sequence
                start_index = inline[i].text.find(key_name[key_index])
                check_length = len(inline[i].text)
                for text_index in range(start_index, check_length):
                    if inline[i].text[text_index] != key_name[key_index]:
                        # no match so must be false positive
                        break
                if key_index == 0:
                    started = True
                chars_found = check_length - start_index
                key_index += chars_found
                found_runs.append((i, start_index, chars_found))
                if key_index != len(key_name):
                    continue
                else:
                    # found all chars in key_name
                    found_all = True
                    break

            # case 2: search for partial text, find subsequent run
            if key_name[key_index] in inline[i].text and started and not found_all:
                # check sequence
                chars_found = 0
                check_length = min(len(inline[i].text), len(key_name[key_index]) - key_index)
                for text_index in range(0, check_length):
                    if inline[i].text[text_index] == key_name[key_index]:
                        key_index += 1
                        chars_found += 1
                    else:
                        break
                # no match so must be end
                found_runs.append((i, 0, chars_found))
                if key_index == len(key_name)

                    found_all = True
                    break

        if found_all and not replace_done:
            for i, item in enumerate(found_runs):
                index, start, length = [t for t in item]
                if i == 0:
                    text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                    inline[index].text = text
                else:
                    text = inline[index].text.replace(inline[index].text[start:start + length], '')
                    inline[index].text = text
            if key_name in p.text:
                replace_inline(key_name, val, inline)
    print('before p')
    for p in paragraphs:
        for key, val in data.items():
            # key_name = '${{{}}}'.format(key)  # use placeholders in the form ${PlaceholderName}
            key_name = key
            print(p.text)
            if key_name in p.text:
                inline = p.runs
                print('asdasd')
                replace_inline(key_name, val, inline)


# usage
le_docx: str = 'example-resumation.docx'

doc = docx.Document(le_docx)
# docx_replace(doc, dict(ItemOne='replacement text', ItemTwo="Some replacement text\nand some more")
docx_replace(doc, dict(ItemOne='dummy', ItemTwo="WOLOL  "))
doc.save(le_docx)