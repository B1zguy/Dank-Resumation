import docx

def shuttle_text(shuttle):
    t = ''
    for i in shuttle:
        t += i.text
    return t

def docx_replace(doc, data):
    for key in data:
        print(key)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if key in cell.text:
                        cell.text = cell.text.replace(key, data[key])

        for p in doc.paragraphs:

            begin = 0
            for end in range(len(p.runs)):

                shuttle = p.runs[begin:end+1]

                full_text = shuttle_text(shuttle)
                if key in full_text:
                    print(key, full_text)
                    # print('Replace：', key, '->', data[key])
                    # print([i.text for i in shuttle])

                    # find the begin
                    index = full_text.index(key)
                    # print('full_text length', len(full_text), 'index:', index)
                    while index >= len(p.runs[begin].text):
                        index -= len(p.runs[begin].text)
                        begin += 1

                    shuttle = p.runs[begin:end+1]

                    # do replace
                    # print('before replace', [i.text for i in shuttle])
                    if key in shuttle[0].text:
                        shuttle[0].text = shuttle[0].text.replace(key, data[key])
                    else:
                        replace_begin_index = shuttle_text(shuttle).index(key)
                        replace_end_index = replace_begin_index + len(key)
                        replace_end_index_in_last_run = replace_end_index - len(shuttle_text(shuttle[:-1]))
                        shuttle[0].text = shuttle[0].text[:replace_begin_index] + data[key]

                        # clear middle runs
                        for i in shuttle[1:-1]:
                            i.text = ''

                        # keep last run
                        shuttle[-1].text = shuttle[-1].text[replace_end_index_in_last_run:]

                    # print('after replace', [i.text for i in shuttle])

                    # set begin to next
                    begin = end

# usage
le_docx: str = 'example-resumation.docx'

doc = docx.Document(le_docx)
# docx_replace(doc, dict(ItemOne='replacement text', ItemTwo="Some replacement text\nand some more")
docx_replace(doc, dict(ItemOne='yeet', ItemTwo="WOLOL  "))
doc.save(le_docx)