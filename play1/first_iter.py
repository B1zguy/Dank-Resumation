from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

le_docx: str = 'example-resumation.docx'
document = Document(le_docx)

for p in document.paragraphs:
    # print(p.text)
    inline = p.runs
    # for i in inline:
    #     print(i.text)
    print('length inline runs per par: ', len(inline))
    for i in range(0, len(inline)):
        if 'yeet' in inline[i].text:
            print(i)
            print(inline[i].text)
            print(len(inline[i].text))
            print('~~~~~~ \n\n')


# document.save(le_docx)



'''
Testing with indentation of paragraaphs

for paragraph in document.paragraphs:
    print(paragraph.paragraph_format.alignment)
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    print(paragraph.text)
    print(paragraph.paragraph_format.alignment)
    print('~~~~~~ \n\n')
    
'''